"""DOCX parser using python-docx."""

from docx import Document as DocxDocument
from app.ingestion.parser import DocumentParser, ParsedDocument, ParsedSection


class DOCXParser(DocumentParser):

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def parse(self, file_path: str, filename: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        sections = []
        tables = []
        current_section = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings as section boundaries
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                # Save previous section
                if current_content:
                    sections.append(ParsedSection(
                        content="\n".join(current_content),
                        section=current_section,
                    ))
                    current_content = []
                current_section = text
                current_content.append(text)
            else:
                current_content.append(text)

        # Save last section
        if current_content:
            sections.append(ParsedSection(
                content="\n".join(current_content),
                section=current_section,
            ))

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                tables.append({
                    "headers": rows[0] if rows else [],
                    "rows": rows[1:] if len(rows) > 1 else [],
                })

        return ParsedDocument(
            sections=sections,
            title=filename,
            metadata={"format": "docx"},
            tables=tables,
        )
